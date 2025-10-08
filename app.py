import os
import zipfile
import tempfile
import gradio as gr
from docx import Document
from huggingface_hub import InferenceClient
# -----------------------------
# 1️⃣ Hugging Face client
# -----------------------------
HF_TOKEN = os.getenv("HUGGINGFACE_API_KEY")
client = InferenceClient(api_key=HF_TOKEN)

# -----------------------------
# 2️⃣ Review prompt
# -----------------------------
promt = """
You are a senior software engineer reviewing .NET / C# code. 
Give short, clear feedback in bullet points.
**Rules:**
Review the code and provide actionable points under these 4 areas:
1. **Code Standards** — Naming, formatting, magic numbers, comments
2. **Security** — SQL injection, input validation, hardcoded secrets, authentication issues
3. **Reusability** — Duplicate code, missing helper functions, not using libraries
4. **Refactoring** — Simplify complex code, performance improvements, better error handling
**Format:**
- Each suggestion must be a single concise line:  
  `Line X: Problem — Fix`
- Always **show both the incorrect and corrected examples** when suggesting naming or syntax improvements.
- Use **real corrected form** (e.g., `_AuthService` → `_authService`).
- Keep tone friendly, direct, and professional.
- Do not repeat identical feedback for multiple lines; combine where possible.
"""

ignore_folders = ['.venv', 'wwwroot', 'node_modules', '__pycache__', 'bin', 'obj', 'properties']
ALLOWED_EXTS = [".py", ".js", ".java", ".cs", ".cpp", ".ts", ".cshtml", ".razor"]

# -----------------------------
# 3️⃣ Helper functions
# -----------------------------
def analyze_code_with_ai(code_text: str, filename: str) -> str:
    try:
        completion = client.chat.completions.create(
            model="meta-llama/Llama-3.1-8B-Instruct",
            messages=[
                {"role": "system", "content": promt},
                {"role": "user", "content": f"Review the following code from {filename}:\n\n{code_text}"}
            ]
        )
        return completion.choices[0].message["content"]
    except Exception as e:
        return f"⚠️ Error: {str(e)}"

def extract_zip_to_temp(zip_file_path):
    temp_dir = tempfile.mkdtemp()
    with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)
    return temp_dir

def list_subfolders(folder_path):
    folders = [os.path.basename(folder_path)]
    for root, dirs, _ in os.walk(folder_path):
        for d in dirs:
            if any(ignored in root for ignored in ignore_folders):
                continue
            rel_path = os.path.relpath(os.path.join(root, d), folder_path)
            folders.append(rel_path)
    return folders

def process_selected_folders(base_path, selected_folders):
    reviews = {}
    for subfolder in selected_folders:
        full_path = base_path if subfolder == os.path.basename(base_path) else os.path.join(base_path, subfolder)
        for root, _, files in os.walk(full_path):
            if any(ignored in root for ignored in ignore_folders):
                continue
            for file in files:
                if any(file.endswith(ext) for ext in ALLOWED_EXTS):
                    filepath = os.path.join(root, file)
                    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                        code = f.read()
                    reviews[file] = analyze_code_with_ai(code, file)
    return reviews

def process_file(file_path):
    filename = os.path.basename(file_path)
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        code = f.read()
    return {filename: analyze_code_with_ai(code, filename)}

def generate_report(reviews, output_path="review_report.docx"):
    doc = Document()
    doc.add_heading("Code Review Report", 0)
    for fname, review in reviews.items():
        doc.add_heading(fname, level=1)
        doc.add_paragraph(review)
    doc.save(output_path)
    return output_path

# -----------------------------
# 4️⃣ Gradio functions
# -----------------------------
# List subfolders after ZIP upload
def load_subfolders_from_zip(zip_file):
    if not zip_file:
        return gr.update(choices=[], value=[]), "⚠️ No ZIP uploaded."
    folder_path = extract_zip_to_temp(zip_file.name)
    folders = list_subfolders(folder_path)
    return gr.update(choices=folders, value=folders), folder_path, f"✅ Found {len(folders)} folders."

# Review selected subfolders
def review_zip_selected(folder_path, selected_folders):
    if not folder_path or not os.path.isdir(folder_path):
        return "⚠️ Invalid folder path.", None
    if not selected_folders:
        return "⚠️ No folders selected.", None
    reviews = process_selected_folders(folder_path, selected_folders)
    if not reviews:
        return "⚠️ No code files found.", None
    report_path = generate_report(reviews)
    output_text = "\n\n".join([f"📌 {f}:\n{r}" for f, r in reviews.items()])
    return output_text, report_path

# Review single file
def review_single_file(file_obj):
    if not file_obj:
        return "⚠️ No file uploaded.", None
    reviews = process_file(file_obj.name)
    if not reviews:
        return "⚠️ Could not analyze file.", None
    report_path = generate_report(reviews)
    output_text = "\n\n".join([f"📌 {f}:\n{r}" for f, r in reviews.items()])
    return output_text, report_path

# -----------------------------
# 5️⃣ Gradio UI
# -----------------------------
if __name__ == "__main__":
 with gr.Blocks() as demo:
    gr.Markdown("# 🤖 AI Code Reviewer\nUpload a ZIP or a single file and get a code review report.")

    with gr.Tab("📦 Review ZIP with Folder Selection"):
     zip_input = gr.File(label="Upload ZIP File", file_types=[".zip"])
     load_btn = gr.Button("📁 Load Subfolders")
     folder_status = gr.Markdown("")
     folder_select = gr.CheckboxGroup(label="Select Folders to Review")
     run_zip_btn = gr.Button("🚀 Run Review")
     zip_output_text = gr.Textbox(label="AI Review Output", lines=15)
     zip_report_file = gr.File(label="Download DOCX Report", type="filepath")

     # Store the temp folder path in gr.State
     temp_folder_state = gr.State()

     # When loading subfolders, also store folder path in state
     load_btn.click(
        fn=load_subfolders_from_zip, 
        inputs=zip_input, 
        outputs=[folder_select, temp_folder_state, folder_status]
    )

    # Use stored folder path from gr.State for review
    run_zip_btn.click(
        fn=review_zip_selected, 
        inputs=[temp_folder_state, folder_select], 
        outputs=[zip_output_text, zip_report_file]
    )

    with gr.Tab("📄 Review Single File"):
        file_input = gr.File(label="Upload Single File", file_types=ALLOWED_EXTS)
        run_file_btn = gr.Button("🚀 Run Review")
        file_output_text = gr.Textbox(label="AI Review Output", lines=15)
        file_report_file = gr.File(label="Download DOCX Report", type="filepath")
        run_file_btn.click(fn=review_single_file, inputs=file_input, outputs=[file_output_text, file_report_file])

demo.launch()
